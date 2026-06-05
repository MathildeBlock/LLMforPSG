import sys
import re
import shutil
import subprocess
import argparse
from pathlib import Path
from tqdm import tqdm

# ---------------------------------------------------------------------------
# KONVERTERING (LibreOffice i stedet for Word)
# ---------------------------------------------------------------------------

def _find_libreoffice():
    """Finder LibreOffice i en VM (Linux eller Windows)."""
    for name in ("soffice", "libreoffice"):
        path = shutil.which(name)
        if path: return path
    # Standard Windows-stier hvis shutil.which fejler
    for candidate in [r"C:\Program Files\LibreOffice\program\soffice.exe", 
                      r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"]:
        if Path(candidate).exists(): return candidate
    return None

def convert_doc_to_docx(doc_path, docx_dir):
    lo_bin = _find_libreoffice()
    if not lo_bin:
        raise RuntimeError("LibreOffice blev ikke fundet. Installer det i din VM (sudo apt install libreoffice).")
    
    # Kør konvertering
    subprocess.run([
        lo_bin, "--headless", "--convert-to", "docx", 
        "--outdir", str(docx_dir), str(doc_path.resolve())
    ], capture_output=True, check=True)
    
    return docx_dir / (doc_path.stem + ".docx")

# ---------------------------------------------------------------------------
# HØJ-PRÆCISIONS UDTRÆKNING (Fra dit originale script)
# ---------------------------------------------------------------------------

_REPLACEMENTS = {
    "\u00b5": "u", "\u00b1": "+/-", "\u2013": "-", "\u2014": "-",
    "\u2019": "'", "\u201c": '"', "\u201d": '"', "\u00b0": " grader",
    "\u2265": ">=", "\u2264": "<=", "\u00d7": "x"
}

def normalize_text(text):
    for char, replacement in _REPLACEMENTS.items():
        text = text.replace(char, replacement)
    text = re.sub(r"[^\S\n]", " ", text) # Non-breaking spaces osv.
    text = re.sub(r"[\x00-\x08\x0b-\x1f\x7f-\x9f]", "", text) # Kontroltegn
    text = re.sub(r" {2,}", " ", text)
    return text.strip()

def is_merge_continuation(tc):
    """Tjekker om cellen er en del af en flettet celle (skal ignoreres)."""
    from docx.oxml.ns import qn
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is not None:
        for tag in ["w:hMerge", "w:vMerge"]:
            merge = tcPr.find(qn(tag))
            if merge is not None and merge.get(qn("w:val"), "restart") != "restart":
                return True
    return False

def extract_table(table):
    lines = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            if not is_merge_continuation(cell._tc):
                txt = normalize_text(cell.text)
                cells.append(txt if txt else "—")
        
        # Rens rækken for ydre tomme markører
        while cells and cells[0] == "—": cells.pop(0)
        while cells and cells[-1] == "—": cells.pop()
        
        line = " | ".join(cells)
        if line and line != "—": lines.append(line)
    return lines

def extract_textboxes(doc):
    from docx.oxml.ns import qn
    texts = []
    # Clark-notation for namespaces der ofte bruges til tekstbokse
    tags = ["{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx",
            "{http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas}txbx",
            "{urn:schemas-microsoft-com:vml}textbox"]
    for tag in tags:
        for txbx in doc.element.body.findall(".//" + tag):
            for p in txbx.findall(".//" + qn("w:p")):
                txt = normalize_text("".join(r.text or "" for r in p.findall(".//" + qn("w:t"))))
                if txt: texts.append(txt)
    return texts

def extract_from_container(container):
    """Hjælpefunktion til at trække tekst ud af headers/footers rækker og tabeller."""
    lines = []
    # Tag alle paragraffer i headeren/footeren
    for p in container.paragraphs:
        txt = normalize_text(p.text)
        if txt:
            lines.append(txt)
    # Tag alle tabeller i headeren/footeren
    for t in container.tables:
        lines.extend(extract_table(t))
    return "\n".join(lines)

def extract_high_fidelity(path):
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    doc = Document(str(path))
    output = []

    # 1. Sidehoved (Rettet her!)
    header_texts = []
    for section in doc.sections:
        # Vi tjekker alle typer headers (normal, første side, lige sider)
        for h_obj in [section.header, section.first_page_header, section.even_page_header]:
            if h_obj:
                txt = extract_from_container(h_obj)
                if txt and txt not in header_texts:
                    header_texts.append(txt)
    
    if header_texts:
        output.append("[SIDEHOVED]\n" + "\n".join(header_texts))

    # 2. Body (Paragraffer og tabeller i rækkefølge)
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            txt = normalize_text(Paragraph(child, doc).text)
            if txt: output.append(txt)
        elif child.tag == qn("w:tbl"):
            output.extend(extract_table(Table(child, doc)))

    # 3. Tekstbokse
    boxes = extract_textboxes(doc)
    if boxes:
        output.append("\n[TEKSTBOKSE]\n" + "\n".join(boxes))

    # 4. Sidefod (Rettet her!)
    footer_texts = []
    for section in doc.sections:
        for f_obj in [section.footer, section.first_page_footer, section.even_page_footer]:
            if f_obj:
                txt = extract_from_container(f_obj)
                if txt and txt not in footer_texts:
                    footer_texts.append(txt)

    if footer_texts:
        output.append("\n[SIDEFOD]\n" + "\n".join(footer_texts))

    return f"=== PSG RAPPORT ===\nFilnavn: {path.name}\n\n" + "\n\n".join(output) + "\n\n=== SLUT ==="


def resolve_input_path(raw_path: str, bdi_reports_dir: Path) -> Path:
    """Resolve a path from the --liste file, preferring paths that actually exist."""

    raw_path = raw_path.strip()
    p = Path(raw_path)
    if not raw_path:
        return p

    def _insert_old_reports_after_bdi(path: Path) -> Path | None:
        parts = list(path.parts)
        parts_lower = [part.lower() for part in parts]
        if "old_reports" in parts_lower:
            return None

        bdi_name_lower = bdi_reports_dir.name.lower()
        try:
            idx = parts_lower.index(bdi_name_lower)
        except ValueError:
            return None

        new_parts = parts[: idx + 1] + ["OLD_reports"] + parts[idx + 1 :]
        return Path(*new_parts)

    candidates: list[Path] = []
    candidates.append(p)

    wants_old_bucket = "old" in p.name.lower()

    # If the list contains relative paths that don't start with bdi_reports, try under it.
    if not p.is_absolute():
        parts_lower = [part.lower() for part in p.parts]
        if not (p.parts and p.parts[0].lower() == bdi_reports_dir.name.lower()):
            candidates.append(bdi_reports_dir / p)
        if wants_old_bucket and "old_reports" not in parts_lower:
            candidates.append(bdi_reports_dir / "OLD_reports" / p)

    # If the path contains bdi_reports anywhere, try inserting OLD_reports after it.
    if wants_old_bucket:
        alt = _insert_old_reports_after_bdi(p)
        if alt is not None:
            candidates.append(alt)

    # Return the first existing candidate; otherwise fall back to the last guess.
    for c in candidates:
        if c.exists():
            return c
    return candidates[-1]

# ---------------------------------------------------------------------------
# MAIN PROCESS
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--mappe", "-m", required=True)
    parser.add_argument("--liste", "-l", required=True)
    parser.add_argument(
        "--bdi-reports-dir",
        default="bdi_reports",
        help="Rodmappe for input-rapporter (bruges hvis --liste indeholder relative stier).",
    )
    parser.add_argument("--max-files", type=int)
    args = parser.parse_args()

    base_dir = Path(args.mappe)
    docx_dir = base_dir / "docx"
    txt_dir = base_dir / "txt"
    docx_dir.mkdir(parents=True, exist_ok=True)
    txt_dir.mkdir(parents=True, exist_ok=True)

    bdi_reports_dir = Path(args.bdi_reports_dir)

    # Indlæs liste
    stier = [
        resolve_input_path(line.split(";")[0].strip(), bdi_reports_dir)
        for line in Path(args.liste).read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]
    if args.max_files: stier = stier[:args.max_files]

    print(f"Starter behandling af {len(stier)} filer...")
    missing_count = 0

    for sti in tqdm(stier):
        if not sti.exists():
            missing_count += 1
            if missing_count <= 10:
                print(f"\n⚠️ Fil ikke fundet (springer over): {sti}")
            continue
        try:
            out_path = txt_dir / (sti.stem + ".txt")
            if out_path.exists():
                continue

            # Trin 1: Konvertér (hvis nødvendigt)
            if sti.suffix.lower() == ".doc":
                docx_path = docx_dir / (sti.stem + ".docx")
                if not docx_path.exists():
                    docx_path = convert_doc_to_docx(sti, docx_dir)
            else:
                # Kopier docx til docx-mappen for at holde det samlet
                docx_path = docx_dir / sti.name
                if not docx_path.exists():
                    shutil.copy2(sti, docx_path)

            # Trin 2: Ekstrahér med høj præcision
            text_content = extract_high_fidelity(docx_path)
            
            # Trin 3: Gem
            out_path.write_text(text_content, encoding="utf-8")
            
        except Exception as e:
            print(f"\n❌ Fejl ved {sti.name}: {e}")

    print(f"\n✅ Færdig! Tekstfiler ligger i: {txt_dir}")
    if missing_count:
        print(f"\nℹ️ Sprang {missing_count} filer over fordi de ikke blev fundet.")

if __name__ == "__main__":
    main()