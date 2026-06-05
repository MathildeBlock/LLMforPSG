import re
import argparse
from pathlib import Path
from tqdm import tqdm

# ---------------------------------------------------------------------------
# TEKSTNORMALISERING
# ---------------------------------------------------------------------------

_REPLACEMENTS = {
    "\u00b5": "u", "\u00b1": "+/-", "\u2013": "-", "\u2014": "-",
    "\u2019": "'", "\u201c": '"',   "\u201d": '"', "\u00b0": " grader",
    "\u2265": ">=", "\u2264": "<=", "\u00d7": "x",
}

def normalize_text(text):
    for char, replacement in _REPLACEMENTS.items():
        text = text.replace(char, replacement)
    text = re.sub(r"[^\S\n]", " ", text)
    text = re.sub(r"[\x00-\x08\x0b-\x1f\x7f-\x9f]", "", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()

# ---------------------------------------------------------------------------
# TABELUDTRÆKNING
# ---------------------------------------------------------------------------

def extract_table(table):
    from docx.oxml.ns import qn
    lines = []
    for row in table.rows:
        cells = []
        for tc in row._tr.tc_lst:
            # Spring vertikalt merged continuation-celler over
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is not None:
                vMerge = tcPr.find(qn("w:vMerge"))
                if vMerge is not None:
                    if vMerge.get(qn("w:val"), "") != "restart":
                        continue

            tekst_dele = []
            for p in tc.findall(".//" + qn("w:p")):
                p_tekst = "".join(
                    r.text or "" for r in p.findall(".//" + qn("w:t"))
                )
                p_tekst = normalize_text(p_tekst)
                if p_tekst:
                    tekst_dele.append(p_tekst)

            cells.append(" ".join(tekst_dele) if tekst_dele else "—")

        while cells and cells[0]  == "—": cells.pop(0)
        while cells and cells[-1] == "—": cells.pop()

        line = " | ".join(cells)
        if line and line != "—":
            lines.append(line)
    return lines

# ---------------------------------------------------------------------------
# UDTRÆKNING
# ---------------------------------------------------------------------------

def extract_from_container(container):
    lines = []
    for p in container.paragraphs:
        txt = normalize_text(p.text)
        if txt:
            lines.append(txt)
    for t in container.tables:
        lines.extend(extract_table(t))
    return "\n".join(lines)

def extract_textboxes(doc):
    from docx.oxml.ns import qn
    texts = []
    tags = [
        "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx",
        "{http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas}txbx",
        "{urn:schemas-microsoft-com:vml}textbox",
    ]
    for tag in tags:
        for txbx in doc.element.body.findall(".//" + tag):
            for p in txbx.findall(".//" + qn("w:p")):
                txt = normalize_text(
                    "".join(r.text or "" for r in p.findall(".//" + qn("w:t")))
                )
                if txt:
                    texts.append(txt)
    return texts

def docx_til_txt(docx_sti: Path) -> str:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    doc = Document(str(docx_sti))
    output = []

    # Sidehoved
    header_texts = []
    for section in doc.sections:
        for h in [section.header, section.first_page_header, section.even_page_header]:
            if h:
                txt = extract_from_container(h)
                if txt and txt not in header_texts:
                    header_texts.append(txt)
    if header_texts:
        output.append("[SIDEHOVED]\n" + "\n".join(header_texts))

    # Body
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            txt = normalize_text(Paragraph(child, doc).text)
            if txt:
                output.append(txt)
        elif child.tag == qn("w:tbl"):
            output.extend(extract_table(Table(child, doc)))

    # Tekstbokse
    boxes = extract_textboxes(doc)
    if boxes:
        output.append("\n[TEKSTBOKSE]\n" + "\n".join(boxes))

    # Sidefod
    footer_texts = []
    for section in doc.sections:
        for f in [section.footer, section.first_page_footer, section.even_page_footer]:
            if f:
                txt = extract_from_container(f)
                if txt and txt not in footer_texts:
                    footer_texts.append(txt)
    if footer_texts:
        output.append("\n[SIDEFOD]\n" + "\n".join(footer_texts))

    return (
        f"=== PSG RAPPORT ===\nFilnavn: {docx_sti.name}\n\n"
        + "\n\n".join(output)
        + "\n\n=== SLUT ==="
    )

# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Konverter DOCX-filer til ren tekst"
    )
    parser.add_argument("--input",  "-i", required=True,
                        help="Mappe med DOCX-filer")
    parser.add_argument("--output", "-o", required=True,
                        help="Mappe til TXT-filer")
    parser.add_argument("--dry-run", action="store_true",
                        help="Vis hvad der ville ske uden at behandle noget")
    parser.add_argument("--overwrite", action="store_true",
                        help="Overskriv eksisterende TXT-filer")
    args = parser.parse_args()

    input_mappe  = Path(args.input)
    output_mappe = Path(args.output)

    if not input_mappe.exists():
        print(f"Fejl: mappen '{input_mappe}' findes ikke.")
        return

    docx_filer = sorted(input_mappe.glob("*.docx"))
    if not docx_filer:
        print(f"Ingen DOCX-filer fundet i '{input_mappe}'.")
        return

    allerede = [f for f in docx_filer
                if (output_mappe / (f.stem + ".txt")).exists()]
    skal     = [f for f in docx_filer
                if args.overwrite or not (output_mappe / (f.stem + ".txt")).exists()]

    print(f"\nFundet:          {len(docx_filer)} DOCX-filer")
    print(f"Allerede gjort:  {len(allerede)}")
    print(f"Skal behandles:  {len(skal)}")

    if args.dry_run:
        print("\n[dry-run] Ingen filer behandlet.")
        return

    if not skal:
        print("\nIntet at gøre.")
        return

    output_mappe.mkdir(parents=True, exist_ok=True)
    fejl = 0

    for docx_sti in tqdm(skal):
        try:
            tekst = docx_til_txt(docx_sti)
            ud = output_mappe / (docx_sti.stem + ".txt")
            ud.write_text(tekst, encoding="utf-8")
        except Exception as e:
            fejl += 1
            print(f"\n {docx_sti.name}: {e}")

    print(f"\n Færdig — tekstfiler i: {output_mappe}")
    if fejl:
        print(f" {fejl} filer fejlede.")

if __name__ == "__main__":
    main()